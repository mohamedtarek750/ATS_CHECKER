"""Generate a set of sample CVs for trying the ATS out.

Every person here is fictional. The set is built to exercise each path through the
pipeline: several roles that should be accepted, two CVs written in an obviously
LLM-generated register, documents that are not CVs at all, and two edge cases.

    python samples/make_samples.py

Then:

    python ats_cli.py --input samples --output data/output --dry-run
"""

from __future__ import annotations

from pathlib import Path

OUT = Path(__file__).resolve().parent
LINES_PER_PAGE = 56


# --------------------------------------------------------------------------
# Minimal PDF writer (no third-party dependency, real extractable text layer)
# --------------------------------------------------------------------------
def _content_stream(lines: list[str]) -> bytes:
    parts = ["BT", "/F1 9 Tf", "1 0 0 1 52 790 Tm", "13 TL"]
    for text in lines:
        size = 15 if text.startswith("# ") else 10 if text.startswith("## ") else 9
        body = text[2:] if text.startswith("# ") else text[3:] if text.startswith("## ") else text
        escaped = body.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        parts.append(f"/F1 {size} Tf")
        parts.append(f"({escaped}) Tj")
        parts.append("T*")
    parts.append("ET")
    return "\n".join(parts).encode("latin-1", errors="replace")


def write_pdf(path: Path, lines: list[str]) -> None:
    pages = [lines[i : i + LINES_PER_PAGE] for i in range(0, len(lines), LINES_PER_PAGE)] or [[]]

    # Object ids: 1 = catalog, 2..1+n = pages, 2+n..1+2n = contents,
    # 2+2n = font, 3+2n = the page tree.
    count = len(pages)
    page_ids = [2 + i for i in range(count)]
    content_ids = [2 + count + i for i in range(count)]
    font_id = 2 + 2 * count
    pages_id = font_id + 1

    objects: list[bytes] = [b""]  # 1: catalog, filled in below

    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    objects[0] = f"<< /Type /Catalog /Pages {pages_id} 0 R >>".encode()

    for index in range(count):
        objects.append(
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 595 842] "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> "
            f"/Contents {content_ids[index]} 0 R >>".encode()
        )
    for page_lines in pages:
        stream = _content_stream(page_lines)
        objects.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
        )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for index, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode()
    path.write_bytes(bytes(out))


def write_docx(path: Path, lines: list[str]) -> None:
    import docx

    document = docx.Document()
    for text in lines:
        if text.startswith("# "):
            document.add_heading(text[2:], level=0)
        elif text.startswith("## "):
            document.add_heading(text[3:], level=1)
        elif text.startswith("- "):
            document.add_paragraph(text[2:], style="List Bullet")
        else:
            document.add_paragraph(text)
    document.save(str(path))


def write_txt(path: Path, lines: list[str]) -> None:
    cleaned = [t[2:] if t.startswith("# ") else t[3:] if t.startswith("## ") else t for t in lines]
    path.write_text("\n".join(cleaned), encoding="utf-8")


# ==========================================================================
# ACCEPT - genuine, human-written CVs across different roles
# ==========================================================================
DATA_ANALYST = """\
# OMAR H. ABDELRAHMAN
Giza, Egypt | +20 111 222 0198 | omar.abdelrahman.data@example.com | linkedin.com/in/omar-abd-data

## SUMMARY
Data analyst, 3 years in retail and FMCG reporting. Most of my day is SQL and Power BI.
Comfortable owning a dashboard end to end, from pulling the data to arguing with the
commercial team about what the number actually means.

## EXPERIENCE
Data Analyst - Alameda Retail Group (Mar 2024 - Present)
- Own the weekly commercial pack: 11 dashboards, ~40 internal users.
- Rewrote the stock-cover query that used to time out. Went from 6 minutes to 40 seconds,
  mostly by fixing a join that was fanning out on returns.
- Built the promo post-mortem template the category managers now use.
- Train new joiners on the warehouse schema. Nobody documented it before me.

Junior Data Analyst - Cairo Foods (Sep 2022 - Feb 2024)
- Daily and monthly sales reporting in Excel, then moved most of it to Power BI.
- Cleaned the customer master. About 12,000 duplicate records from three legacy systems.
- Supported the finance close with ad-hoc extracts.

Intern - Cairo Foods (Jul 2022 - Aug 2022)
- Data entry and validation for the supplier catalogue.

## EDUCATION
BSc Statistics - Cairo University, 2022. Graduation project on demand forecasting for
a small bakery chain, using SARIMA. Grade: Very Good.

## SKILLS
SQL (advanced - window functions, CTEs, query tuning), Power BI / DAX, Excel (Power Query),
Python (pandas, matplotlib - intermediate), Git basics, some Airflow exposure.

## CERTIFICATIONS
Microsoft PL-300 Power BI Data Analyst (2024)
Google Data Analytics Certificate (2023)

## LANGUAGES
Arabic (native), English (fluent)
"""

FRONTEND_DEV = """\
# NOURAN K. SAAD
Alexandria, Egypt | +20 106 774 3311 | nouran.saad.dev@example.com
github.com/nouran-saad | nouransaad.dev

## ABOUT
Front-end developer, 4 years. React and TypeScript mostly. I care a lot about
accessibility and I am the person on the team who actually runs the screen reader.

## WORK
Front-end Developer - Bokra Technologies (Jan 2023 - Present)
- Main developer on the customer portal. React 18, TypeScript, Vite, TanStack Query.
- Took the Lighthouse performance score from 42 to 89, mostly route-level code splitting
  and killing a 400kb date library we used in two places.
- Migrated the component library from styled-components to Tailwind over about 4 months.
  Still three legacy screens left, honestly.
- Set up the a11y lint rules and fixed roughly 200 violations. Got the portal through
  the client's WCAG AA audit.

Front-end Developer - Freelance (Jun 2021 - Dec 2022)
- Built and maintained six small business sites. Next.js and WordPress headless.
- One e-commerce front-end for a local furniture shop, integrated with their existing
  PHP backend which was... an experience.

## EDUCATION
BSc Computer Science - Alexandria University, 2021

## SKILLS
React, TypeScript, JavaScript (ES2022), Next.js, Tailwind CSS, HTML, CSS,
TanStack Query, Zustand, Vitest, Playwright, Figma handoff, Git, basic Node.js

## THINGS I HAVE WRITTEN
"Why your focus ring matters" - talk at Alex Frontend Meetup, Nov 2024
"""

BACKEND_DEV = """\
# YOUSSEF M. TAHA
Nasr City, Cairo | +20 122 908 4417 | youssef.taha.be@example.com

## PROFILE
Backend engineer with 5 years building APIs, mostly Python and some Go. Last two years
have been heavy on payments, which means a lot of time thinking about idempotency.

## EXPERIENCE
Senior Backend Engineer - PayLink MEA (Nov 2023 - Present)
- Own the settlement service. Django + Celery + Postgres, roughly 300k transactions/day.
- Redesigned the retry logic after an incident where a provider timeout caused double
  settlement on 84 merchant accounts. Post-mortem is internal but I wrote it.
- Moved reporting reads onto a replica, which fixed the nightly lock contention.
- On call one week in four.

Backend Engineer - Sahl Systems (Aug 2020 - Oct 2023)
- REST APIs for a logistics platform. FastAPI, Postgres, Redis.
- Wrote the driver-assignment service in Go because the Python version could not keep
  up with the geospatial queries.
- Introduced Alembic migrations. Before that people ran SQL by hand in production.

## EDUCATION
BSc Computer Engineering - Ain Shams University, 2020

## TECHNICAL
Languages: Python, Go, SQL, some Bash
Frameworks: Django, DRF, FastAPI, Celery
Data: PostgreSQL, Redis, Elasticsearch (basic)
Infra: Docker, GitHub Actions, AWS (ECS, RDS, S3), Terraform (can read it, not write it)
Other: pytest, OpenAPI, Sentry

## LANGUAGES
Arabic (native), English (professional)
"""

CIVIL_ENGINEER = """\
# HASSAN A. EL-SHERBINY
New Administrative Capital, Egypt | +20 100 448 2276 | hassan.sherbiny.eng@example.com

## SUMMARY
Site civil engineer, 7 years on residential and mixed-use projects. Three of those years
in the New Capital. Comfortable with reinforced concrete execution and subcontractor
coordination.

## EXPERIENCE
Site Engineer - Orascom Construction (Feb 2022 - Present)
- Execution of two residential towers, R+12, in the New Capital R7 district.
- Supervise concrete pours, steel fixing inspection, and formwork release.
- Coordinate with the consultant on shop drawing approvals. Currently about 40 open items.
- Track quantities and prepare monthly progress invoices for the client.

Site Engineer - Hassan Allam Construction (Sep 2018 - Jan 2022)
- Infrastructure package: roads, storm drainage, and utility trenches for a compound
  in 6th of October.
- Managed a crew of about 30 including three subcontractors.
- Prepared as-built drawings on AutoCAD.

## EDUCATION
BSc Civil Engineering - Cairo University, Faculty of Engineering, 2018
Structural Engineering department. Graduation project: 15-storey RC building design.

## SKILLS
AutoCAD, Revit (intermediate), Primavera P6, MS Project, SAP2000, Safe,
Quantity surveying, ECP 203 code, Site safety, BOQ preparation

## CERTIFICATIONS
OSHA 30-Hour Construction Safety (2023)
PMP - in progress, exam booked for Nov 2026
"""

ML_ENGINEER = """\
# SALMA R. GHONEIM
Maadi, Cairo | +20 128 335 7790 | salma.ghoneim.ml@example.com | github.com/salma-ghoneim

## SUMMARY
ML engineer, 4 years. I spend more time on serving and monitoring than on modelling,
which I think is the correct ratio.

## EXPERIENCE
Machine Learning Engineer - Rakiza AI (Jun 2023 - Present)
- Own the Arabic speech-to-text pipeline. Fine-tuned Whisper-large on ~900 hours of
  Egyptian dialect audio. WER went from 31% to 19% on our internal test set, which is
  still not good enough for the medical use case.
- Built the serving stack: Triton, ONNX exports, autoscaling on GKE.
- Set up drift monitoring after we shipped a model that quietly degraded for six weeks
  before anyone noticed.

Data Scientist - Rakiza AI (Sep 2022 - May 2023)
- Churn model for the B2B product. Gradient boosting, nothing exotic. The feature that
  mattered most turned out to be support ticket volume.

## EDUCATION
MSc Computer Science (Machine Learning) - Nile University, 2022
Thesis on low-resource Arabic ASR.
BSc Computer Science - Nile University, 2020

## SKILLS
Python, PyTorch, HuggingFace Transformers, scikit-learn, ONNX, Triton Inference Server,
MLflow, Docker, Kubernetes, GCP (Vertex, GKE), SQL, Weights & Biases, FastAPI

## PUBLICATIONS
Ghoneim, S. et al. "Dialectal variation in Egyptian Arabic ASR", ArabicNLP Workshop, 2023
"""

UIUX_DESIGNER = """\
# LAILA M. ZAKI
Cairo, Egypt | +20 109 662 1145 | laila.zaki.design@example.com | behance.net/lailazaki

## ABOUT ME
Product designer, 5 years. Fintech for the last three. I run the design system and I do
my own research, which is not ideal but it is where we are.

## EXPERIENCE
Senior Product Designer - Fawry Digital (Apr 2023 - Present)
- Redesigned the merchant onboarding flow. Drop-off between step 2 and 3 fell from
  38% to 21% after we split the KYC upload into its own screen.
- Maintain the design system in Figma. 140+ components, used by 4 squads.
- Ran 22 moderated usability sessions last year, mostly in Arabic, mostly on WhatsApp
  video because merchants would not install anything.

Product Designer - Vodafone Egypt (Jan 2021 - Mar 2023)
- Worked on the self-service app. Bill payment and plan upgrade journeys.
- Accessibility audit and remediation for the Arabic RTL layouts.

## EDUCATION
BA Graphic Design - Helwan University, Faculty of Applied Arts, 2020

## TOOLS
Figma (expert), FigJam, Maze, Hotjar, Adobe Illustrator, After Effects (basic),
HTML/CSS reading level, Notion

## LANGUAGES
Arabic (native), English (fluent), French (basic)
"""

# ==========================================================================
# REJECT - written in an unmistakably LLM-generated register
# ==========================================================================
AI_DATA_SCIENTIST = """\
# ALEXANDER J. MITCHELL
Senior Data Scientist | alexander.mitchell@email.com | +1 (555) 847-2910
LinkedIn Profile | GitHub Profile | Portfolio Website

## PROFESSIONAL SUMMARY
Results-driven Data Scientist with a proven track record of leveraging cutting-edge machine
learning solutions to drive transformative business outcomes. Passionate about leveraging
robust, scalable AI systems to deliver measurable impact across diverse stakeholder
ecosystems. Adept at translating complex analytical insights into actionable strategic
recommendations that empower data-driven decision-making at scale.

## PROFESSIONAL EXPERIENCE
Senior Data Scientist | XYZ Company | 2021 - Present
- Spearheaded the development of predictive models, increasing revenue by 40%.
- Leveraged advanced analytics frameworks, reducing operational costs by 30%.
- Utilized cutting-edge deep learning architectures, improving model accuracy by 25%.
- Orchestrated cross-functional collaboration, accelerating project delivery by 35%.
- Pioneered innovative data strategies, enhancing stakeholder satisfaction by 45%.

Data Scientist | ABC Corporation | 2018 - 2021
- Spearheaded end-to-end data pipelines, enhancing throughput by 50%.
- Leveraged statistical methodologies, optimizing decision-making processes by 45%.
- Utilized robust visualization frameworks, increasing dashboard adoption by 60%.
- Orchestrated model deployment workflows, reducing inference latency by 20%.
- Championed data governance initiatives, improving data quality by 55%.

Junior Data Scientist | Tech Solutions Inc | 2016 - 2018
- Spearheaded exploratory data analysis, uncovering insights that boosted KPIs by 30%.
- Leveraged machine learning algorithms, enhancing classification performance by 25%.
- Utilized cloud infrastructure, decreasing computational overhead by 40%.

## EDUCATION
Master of Science in Data Science | Prestigious University | 2016
Bachelor of Science in Computer Science | Prestigious University | 2014

## TECHNICAL SKILLS
Python, R, SQL, Scala, Java, TensorFlow, PyTorch, Keras, Scikit-learn, XGBoost, LightGBM,
Spark, Hadoop, Hive, Kafka, Airflow, AWS, Azure, GCP, Docker, Kubernetes, Tableau,
Power BI, Looker, Git, Jenkins, Snowflake, Databricks, MLflow, Kubeflow

## KEY ACHIEVEMENTS
- Recognized as a top performer for three consecutive years.
- Delivered transformative solutions that generated significant business value.
- Mentored junior team members, fostering a culture of continuous improvement.
"""

AI_FULLSTACK_TEMPLATE = """\
# [YOUR NAME]
Full Stack Developer | [your.email@example.com] | [Phone Number] | [City, Country]
[LinkedIn URL] | [GitHub URL]

## PROFESSIONAL SUMMARY
Certainly! Here is a professional summary for a Full Stack Developer:
Dynamic and detail-oriented Full Stack Developer with [X] years of experience architecting
end-to-end web solutions. Proven ability to leverage modern frameworks to deliver seamless,
scalable, and robust applications that drive user engagement and business growth.

## PROFESSIONAL EXPERIENCE
Senior Full Stack Developer | [Company Name] | [Start Date] - Present
- Spearheaded the development of scalable web applications, improving performance by 40%.
- Leveraged React and Node.js to deliver seamless user experiences, increasing engagement by 35%.
- Utilized robust CI/CD pipelines, reducing deployment time by 50%.
- Collaborated with cross-functional teams to deliver cutting-edge solutions on schedule.

Full Stack Developer | XYZ Company | [Start Date] - [End Date]
- Spearheaded API development, enhancing system throughput by 45%.
- Leveraged cloud-native architectures, reducing infrastructure costs by 30%.
- Utilized agile methodologies, accelerating sprint velocity by 25%.

## EDUCATION
[Degree] in [Field of Study] | [University Name] | [Graduation Year]

## TECHNICAL SKILLS
Frontend: React, Angular, Vue.js, HTML5, CSS3, JavaScript, TypeScript
Backend: Node.js, Express, Django, Flask, Spring Boot
Databases: MySQL, PostgreSQL, MongoDB, Redis
DevOps: Docker, Kubernetes, Jenkins, AWS, Azure, GCP

Would you like me to tailor this resume for a specific job description?
"""

# ==========================================================================
# REJECT - not a CV at all
# ==========================================================================
COVER_LETTER = """\
Mariam Abdelaziz
15 El Nasr Street, Nasr City, Cairo
mariam.abdelaziz.apply@example.com

25 August 2026

Hiring Manager
Administrative Capital for Urban Development
New Administrative Capital, Egypt

Dear Hiring Manager,

I am writing to apply for the Data Analyst position advertised on your careers page last
week. I graduated from Ain Shams University in 2024 with a degree in Economics and I have
spent the last eighteen months doing reporting work for a mid-size distribution company.

What draws me to ACUD specifically is the scale of the data problem. A city being built
from nothing generates a kind of dataset that does not exist anywhere else in Egypt, and I
would very much like to work on it rather than on quarterly sales variance for the rest of
my twenties.

I have attached my CV, which goes into more detail on the reporting stack I have used and
the two dashboards I built end to end. I am available for an interview at any time and can
start with two weeks notice.

Thank you for your consideration.

Yours sincerely,

Mariam Abdelaziz
"""

CERTIFICATE = """\
CERTIFICATE OF COMPLETION

This is to certify that

KARIM ADEL MOUSTAFA

has successfully completed the course

INTRODUCTION TO MACHINE LEARNING WITH PYTHON
40 contact hours

conducted by the Information Technology Institute (ITI)
between 3 February 2026 and 14 March 2026

with an overall grade of: Excellent (92%)

Certificate ID: ITI-ML-2026-08841
Verify at: iti.gov.eg/verify

________________________          ________________________
Course Instructor                 Programme Director
"""

JOB_POSTING = """\
JOB POSTING - Senior Front-End Developer

Department: Digital Services
Location: New Administrative Capital (on-site, 5 days)
Employment type: Full time, permanent
Posted: 18 August 2026
Closes: 15 September 2026

ABOUT THE ROLE
We are looking for a Senior Front-End Developer to join the Digital Services team working
on citizen-facing portals for the new capital.

RESPONSIBILITIES
- Build and maintain React applications used by hundreds of thousands of residents.
- Own front-end architecture decisions and mentor two mid-level developers.
- Work with the design system team to keep the component library consistent.
- Ensure WCAG 2.1 AA compliance across all delivered screens.

REQUIREMENTS
- 5+ years of professional front-end experience.
- Deep React and TypeScript knowledge.
- Experience with RTL / Arabic layouts is strongly preferred.
- Bachelor degree in Computer Science or equivalent practical experience.

WHAT WE OFFER
Competitive salary, medical insurance for you and dependents, transport allowance,
and an annual training budget.

To apply, send your CV to careers@example-acud.eg quoting reference FE-2026-114.
"""

# ==========================================================================
# EDGE CASES
# ==========================================================================
TOO_SHORT = """\
Ahmed Samir
Looking for a job
ahmed@example.com
"""

AMBIGUOUS_ROLE = """\
# TAREK N. FOUAD
Cairo, Egypt | +20 115 330 8876 | tarek.fouad.work@example.com

## SUMMARY
Fresh graduate. I have done a bit of everything and I am not completely sure what I want
to specialise in yet, but I like building things that people actually open.

## EDUCATION
BSc Computer Science - Modern Academy, Cairo, 2026. GPA 3.1.

## PROJECTS
Graduation project - hospital appointment system
  React front-end, Node/Express API, MySQL. I did the front-end and half the API.
Kaggle Titanic and House Prices
  Followed two tutorials, got a reasonable score, understood maybe 60% of it.
Instagram clone
  Flutter. Never finished the messaging part.
Small Arabic sentiment classifier
  scikit-learn, TF-IDF and logistic regression on a dataset from a course.

## SKILLS
JavaScript, React, Node.js, Python, pandas, scikit-learn, Flutter, Dart, MySQL,
HTML, CSS, Git

## OTHER
Volunteered as a session organiser for the university IEEE branch.
Arabic (native), English (good)
"""


SAMPLES: list[tuple[str, str]] = [
    # accepted - real CVs, varied roles
    ("01_data_analyst_omar.pdf", DATA_ANALYST),
    ("02_frontend_nouran.pdf", FRONTEND_DEV),
    ("03_backend_youssef.docx", BACKEND_DEV),
    ("04_civil_engineer_hassan.pdf", CIVIL_ENGINEER),
    ("05_ml_engineer_salma.docx", ML_ENGINEER),
    ("06_uiux_designer_laila.pdf", UIUX_DESIGNER),
    # rejected - AI generated
    ("07_ai_generated_data_scientist.pdf", AI_DATA_SCIENTIST),
    ("08_ai_generated_fullstack_template.docx", AI_FULLSTACK_TEMPLATE),
    # rejected - not a CV
    ("09_cover_letter.pdf", COVER_LETTER),
    ("10_certificate.txt", CERTIFICATE),
    ("11_job_posting.txt", JOB_POSTING),
    # edge cases
    ("12_too_short.txt", TOO_SHORT),
    ("13_ambiguous_role.pdf", AMBIGUOUS_ROLE),
]


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for name, body in SAMPLES:
        lines = body.rstrip().split("\n")
        path = OUT / name
        if name.endswith(".pdf"):
            write_pdf(path, lines)
        elif name.endswith(".docx"):
            write_docx(path, lines)
        else:
            write_txt(path, lines)
        print(f"  wrote {name:<44} {path.stat().st_size:>7,} bytes")
    print(f"\n{len(SAMPLES)} sample files in {OUT}")


if __name__ == "__main__":
    main()
